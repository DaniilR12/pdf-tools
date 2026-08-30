import sys
from io import BytesIO

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt

source_path, target_path = sys.argv[1:3]
pdf = fitz.open(source_path)
document = Document()

for page_index, page in enumerate(pdf):
    if page_index:
        section = document.add_section(WD_SECTION.NEW_PAGE)
    else:
        section = document.sections[0]

    width = page.rect.width / 72
    height = page.rect.height / 72
    section.page_width = Inches(width)
    section.page_height = Inches(height)
    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    pixels = page.get_pixmap(matrix=fitz.Matrix(2, 2), alpha=False)
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1
    paragraph.paragraph_format.first_line_indent = Inches(0)
    paragraph.add_run().add_picture(BytesIO(pixels.tobytes("png")), width=Inches(width))

pdf.close()
document.save(target_path)
