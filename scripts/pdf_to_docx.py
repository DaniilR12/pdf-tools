import os
import sys
from io import BytesIO

import fitz
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter
from docxcompose.composer import Composer


MATH_FONT_HINTS = (
    "cambria math",
    "stix",
    "asana math",
    "latin modern math",
    "cmsy",
    "cmmi",
    "cmex",
    "msam",
    "msbm",
)

MATH_UNICODE_RANGES = [
    (0x2200, 0x22FF),    # Mathematical Operators
    (0x27C0, 0x27EF),    # Misc Math Symbols-A
    (0x2980, 0x29FF),    # Misc Math Symbols-B
    (0x2A00, 0x2AFF),    # Supplemental Math Operators
    (0x1D400, 0x1D7FF),  # Math Alphanumeric Symbols
]


def has_math_content(page):
    """
    Detect pages containing mathematical notation that pdf2docx
    is unlikely to reconstruct correctly.
    """

    text_dict = page.get_text("dict")

    math_char_count = 0
    total_char_count = 0

    for block in text_dict.get("blocks", []):
        if block.get("type") != 0:
            continue

        for line in block.get("lines", []):
            for span in line.get("spans", []):
                text = span.get("text", "")
                font_name = (span.get("font") or "").lower()

                total_char_count += len(text)

                if any(hint in font_name for hint in MATH_FONT_HINTS):
                    math_char_count += len(text)

                for char in text:
                    codepoint = ord(char)

                    if any(
                        start <= codepoint <= end
                        for start, end in MATH_UNICODE_RANGES
                    ):
                        math_char_count += 1

    if math_char_count == 0:
        return False

    # Do not convert an entire page to an image just because
    # there is one mathematical symbol somewhere.
    #
    # Only treat it as a math-heavy page when math content
    # represents a meaningful portion of the page.
    if total_char_count == 0:
        return False

    return math_char_count >= max(10, total_char_count * 0.12)


def is_scanned_page(page):
    """
    A scanned PDF page usually has no text layer at all.

    Such a page cannot be converted into editable text without OCR.
    In that case, rendering it as an image is preferable to
    producing an empty DOCX page.
    """

    text = page.get_text("text").strip()

    if text:
        return False

    # Check whether the page contains images.
    images = page.get_images(full=True)

    return len(images) > 0


def classify_page(page):
    """
    True  -> render page as image
    False -> convert page using pdf2docx
    """

    # Scanned/image-only PDF
    if is_scanned_page(page):
        return True

    # Math-heavy pages are difficult to reconstruct reliably.
    if has_math_content(page):
        return True

    # IMPORTANT:
    # Do NOT classify multi-column layouts as images.
    #
    # Resumes, CVs, brochures and modern documents often use
    # two-column layouts. pdf2docx should be allowed to attempt
    # converting them into editable content.
    return False


def render_page_to_docx(page, output_path):
    """
    Render one PDF page into a DOCX containing the page as an image.

    Used only for scanned/image-only pages or math-heavy pages.
    """

    document = Document()

    section = document.sections[0]

    width_in = page.rect.width / 72
    height_in = page.rect.height / 72

    section.page_width = Inches(width_in)
    section.page_height = Inches(height_in)

    section.top_margin = Inches(0)
    section.bottom_margin = Inches(0)
    section.left_margin = Inches(0)
    section.right_margin = Inches(0)

    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)

    pixmap = page.get_pixmap(
        matrix=fitz.Matrix(2, 2),
        alpha=False,
    )

    paragraph = document.add_paragraph()

    paragraph.paragraph_format.space_before = 0
    paragraph.paragraph_format.space_after = 0
    paragraph.paragraph_format.line_spacing = 1

    paragraph.add_run().add_picture(
        BytesIO(pixmap.tobytes("png")),
        width=Inches(width_in),
    )

    document.save(output_path)


def convert_pages_with_pdf2docx(
    source_path,
    page_numbers,
    output_path,
):
    """
    Convert a contiguous group of pages using pdf2docx.
    """

    converter = Converter(source_path)

    try:
        converter.convert(
            output_path,

            pages=page_numbers,

            multi_processing=False,

            # Tables
            parse_lattice_table=True,
            parse_stream_table=True,

            connected_border_tolerance=1.5,
            min_border_clearance=1.0,
            max_border_width=8.0,

            # Text layout
            line_break_width_ratio=0.3,
            line_break_free_space_ratio=0.08,
            line_separate_threshold=3.0,
            new_paragraph_free_space_ratio=0.92,

            # Shapes / images
            shape_min_dimension=0.5,

            min_svg_gap_dx=2.0,
            min_svg_gap_dy=0.5,
            min_svg_w=0.5,
            min_svg_h=0.5,

            clip_image_res_ratio=1.0,
            float_image_ignorable_gap=2.0,

            # Page margins
            page_margin_factor_top=0.5,
            page_margin_factor_bottom=0.5,
        )

    finally:
        converter.close()


def group_contiguous_runs(indices):
    """
    [0, 1, 2, 5, 6] -> [[0, 1, 2], [5, 6]]
    """

    runs = []

    for index in indices:
        if runs and index == runs[-1][-1] + 1:
            runs[-1].append(index)
        else:
            runs.append([index])

    return runs


def create_empty_docx(target_path):
    document = Document()
    document.save(target_path)


def main():
    if len(sys.argv) < 3:
        print(
            "Usage: python pdf_to_docx.py <source.pdf> <target.docx>",
            file=sys.stderr,
        )
        sys.exit(1)

    source_path = sys.argv[1]
    target_path = sys.argv[2]

    if not os.path.exists(source_path):
        print(
            f"Source PDF does not exist: {source_path}",
            file=sys.stderr,
        )
        sys.exit(1)

    # ---------------------------------------------------------
    # First pass:
    # inspect every page
    # ---------------------------------------------------------

    pdf = fitz.open(source_path)

    page_count = pdf.page_count

    if page_count == 0:
        pdf.close()
        create_empty_docx(target_path)
        return

    page_is_image = [
        classify_page(pdf[index])
        for index in range(page_count)
    ]

    pdf.close()

    # ---------------------------------------------------------
    # Create fragments
    # ---------------------------------------------------------

    fragment_of_page = {}
    fragments_to_clean = set()

    # Pages that should remain editable
    editable_indices = [
        index
        for index, is_image in enumerate(page_is_image)
        if not is_image
    ]

    # Convert contiguous editable page ranges together.
    #
    # This is important because pdf2docx can preserve the
    # relationship between neighboring pages better than
    # converting every page completely independently.
    for run in group_contiguous_runs(editable_indices):
        fragment_path = (
            f"{target_path}.frag.{run[0]}.docx"
        )

        try:
            convert_pages_with_pdf2docx(
                source_path,
                run,
                fragment_path,
            )

        except Exception as error:
            print(
                f"pdf2docx failed for pages {run}: {error}",
                file=sys.stderr,
            )

            # Do not silently turn a normal editable page into
            # an image just because pdf2docx encountered a problem.
            #
            # Re-raise so the API can report the conversion failure.
            raise

        fragments_to_clean.add(fragment_path)

        for index in run:
            fragment_of_page[index] = fragment_path

    # ---------------------------------------------------------
    # Image fallback pages
    # ---------------------------------------------------------

    pdf = fitz.open(source_path)

    for index, is_image in enumerate(page_is_image):
        if not is_image:
            continue

        fragment_path = (
            f"{target_path}.frag.{index}.docx"
        )

        render_page_to_docx(
            pdf[index],
            fragment_path,
        )

        fragments_to_clean.add(fragment_path)
        fragment_of_page[index] = fragment_path

    pdf.close()

    # ---------------------------------------------------------
    # Merge fragments while preserving page order
    # ---------------------------------------------------------

    master_document = None
    composer = None

    last_fragment_path = None

    for index in range(page_count):
        fragment_path = fragment_of_page[index]

        # A pdf2docx fragment may contain several consecutive
        # pages. Do not append it multiple times.
        if fragment_path == last_fragment_path:
            continue

        if master_document is None:
            master_document = Document(fragment_path)
            composer = Composer(master_document)
        else:
            composer.append(
                Document(fragment_path)
            )

        last_fragment_path = fragment_path

    if composer is None:
        create_empty_docx(target_path)
    else:
        composer.save(target_path)

    # ---------------------------------------------------------
    # Cleanup temporary fragments
    # ---------------------------------------------------------

    for fragment_path in fragments_to_clean:
        try:
            os.remove(fragment_path)
        except OSError:
            pass


if __name__ == "__main__":
    main()